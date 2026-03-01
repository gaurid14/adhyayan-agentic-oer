import asyncio
import io
import json
import os
import tempfile
import threading
import urllib
import re
from typing import List, Dict
from django.urls import reverse
from django.conf import settings
from django.contrib.admin.utils import unquote
from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse, HttpResponseBadRequest
from django.contrib import messages
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt

from PyPDF2 import PdfReader
from docx import Document
from exceptiongroup import ExceptionGroup
from google.auth.exceptions import RefreshError
from langchain_core.messages import HumanMessage
from langsmith import traceable
from xhtml2pdf import pisa
from googleapiclient.http import (
    MediaFileUpload,
    MediaIoBaseUpload,
    MediaIoBaseDownload
)

from accounts.models import (
    Chapter, UploadCheck, Assessment,
    Question, Option, Course, User, ExternalResource, ChapterContributionProgress, ChapterPolicy, AssessmentSource
)
from accounts.views.email.email_service import ContributionSuccessEmail
from langgraph_agents.review_graph.review_runner import run_review_pipeline

from langgraph_agents.services.drive_service import (
    GoogleDriveAuthService,
    GoogleDriveFolderService
)

from langgraph_agents.agents.submission_agent import submission_agent
from langgraph_agents.graph.workflow import compiled_graph
from langgraph_agents.services.evaluation_score import finalize_evaluation
from langgraph_agents.services.gemini_service import llm

import traceback
import sys

from asgiref.sync import sync_to_async
from mcp.client.session import ClientSession
from mcp.client.stdio import stdio_client, StdioServerParameters

from accounts.models import ContentCheck


# PROJECT_ROOT = r"C:\Users\gauri\IdeaProjects\oer"
# MCP_PATH = os.path.join(PROJECT_ROOT, "langgraph_agents", "services", "mcp_server.py")
# ✅ ALWAYS point to project root (folder that has manage.py)
PROJECT_ROOT = str(settings.BASE_DIR)

# MCP server path from project root
MCP_PATH = os.path.join(PROJECT_ROOT, "langgraph_agents", "services", "mcp_server.py")

def get_chapter_storage_usage(service, folder_id):

    total_size = 0

    results = service.files().list(
        q=f"'{folder_id}' in parents and trashed=false",
        fields="files(size)"
    ).execute()

    for f in results.get("files", []):
        total_size += int(f.get("size", 0))

    return total_size


class ContributorSessionService:
    @staticmethod
    def store_submission_context(request, **kwargs):
        request.session.update(kwargs)

    @staticmethod
    def get(request, key, default=None):
        return request.session.get(key, default)



class ContributorSubmissionService:
    @staticmethod
    def has_existing_submission(contributor_id, chapter_id) -> bool:
        return UploadCheck.objects.filter(
            contributor_id=contributor_id,
            chapter_id=chapter_id
        ).exists()


class ContributorDriveUploadService:
    def __init__(self):
        self.service = GoogleDriveAuthService.get_service()
        self.folder_service = GoogleDriveFolderService(self.service)
        self.oer_root_id = self.folder_service.get_or_create_folder("oer_content")

    def ensure_topic_folder(self, base_folder, folder_type, topic=None):
        root = self.folder_service.get_or_create_folder(
            settings.GOOGLE_DRIVE_FOLDERS[folder_type],
            self.oer_root_id
        )
        contributor_folder = self.folder_service.get_or_create_folder(
            base_folder, root
        )
        if topic:
            topic = topic.replace("/", "_").strip()
            return self.folder_service.get_or_create_folder(topic, contributor_folder)
        return contributor_folder

    def upload_file_bytes(self, file_bytes, filename, folder_id, content_type, contributor_id):
        buffer = io.BytesIO(file_bytes)
        media = MediaIoBaseUpload(
            buffer,
            mimetype=content_type,
            resumable=True
        )

        file_metadata = {
            "name": filename,
            "parents": [folder_id]
        }

        self.service.files().create(
            body=file_metadata,
            media_body=media,
            fields="id"
        ).execute()

        # after upload
        buffer.close()


class ContributorEditorService:
    def __init__(self):
        self.service = GoogleDriveAuthService.get_service()
        self.folder_service = GoogleDriveFolderService(self.service)
        self.oer_root = self.folder_service.get_or_create_folder("oer_content")

    def save_draft(self, content, filename, folder_id):
        doc = Document()
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, "html.parser")
        for p in soup.find_all(["p", "div"]):
            if p.text.strip():
                doc.add_paragraph(p.text.strip())

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)

        media = MediaIoBaseUpload(
            bio,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        self.service.files().create(
            body={'name': filename, 'parents': [folder_id]},
            media_body=media,
            fields="id"
        ).execute()


class SubmissionOrchestrator:
    @staticmethod
    @traceable(name="Submission + Evaluation Orchestrator")
    def submit_and_evaluate(state: dict):

        # 1) submission agent sync safe
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            result = loop.run_until_complete(
                submission_agent.ainvoke({
                    "contributor_id": state["contributor_id"],
                    "chapter_id": state["chapter_id"],
                    "drive_folders": state["drive_folders"],
                })
            )
            print("Submission agent invoked!!")
        finally:
            loop.close()

        if result.get("status") != "success":
            return result

        upload_id = result["upload_id"]

        # 2) evaluation graph in background thread
        def run_graph_background():
            bg_loop = asyncio.new_event_loop()
            asyncio.set_event_loop(bg_loop)

            async def wait_for_extraction():
                max_wait_seconds = 300
                interval = 5
                waited = 0

                while waited < max_wait_seconds:
                    status = await sync_to_async(
                        lambda: ContentCheck.objects.filter(upload_id=upload_id)
                        .values_list("extraction_status", flat=True)
                        .first()
                    )()

                    print("🔎 extraction_status =", status)

                    if status is True:
                        print("Extraction confirmed. Starting evaluation graph...")
                        return True

                    await asyncio.sleep(interval)
                    waited += interval

                print("Extraction timeout. Evaluation graph NOT started.")
                return False

            @traceable(name="Adhyayan Evaluation Pipeline")
            async def runner():
                ok = await wait_for_extraction()
                if not ok:
                    return

                assert os.path.exists(MCP_PATH), f"MCP_PATH missing: {MCP_PATH}"

                server_params = StdioServerParameters(
                    command=sys.executable,
                    args=[MCP_PATH],
                    env={
                        **os.environ,
                        "DJANGO_SETTINGS_MODULE": "oer.settings",
                        "PYTHONPATH": PROJECT_ROOT,
                    }
                )

                # start MCP session ONCE
                async with stdio_client(server_params) as (read, write):
                    async with ClientSession(read, write) as session:
                        await session.initialize()
                        print("MCP session initialized")

                        graph_input = {**state, **result}
                        graph_input["mcp_session"] = session   # shared for all agents

                        await compiled_graph.ainvoke(
                            graph_input,
                            config={
                                "run_name": "Adhyayan Evaluation Graph"
                            }
                        )

                        await sync_to_async(UploadCheck.objects.filter(id=upload_id).update)(evaluation_status=True)

                        # writing scores to blockchain
                        await finalize_evaluation(upload_id)
                        print(f"✅ Marked upload {upload_id} as evaluated.")    

                print("Evaluation graph invoked!!")

                # 🔔 Auto-trigger Decision Maker (event-driven; no polling)
                # Runs only if: (1) deadline is over AND (2) min contributions met.
                try:
                    from accounts.services.auto_decision import trigger_decision_if_due  # local import avoids cycles
                    chapter_id = int(state["chapter_id"])
                    dm_run = await sync_to_async(trigger_decision_if_due)(chapter_id)
                    if dm_run:
                        status = getattr(dm_run, "status", None) or (dm_run.get("status") if isinstance(dm_run, dict) else "unknown")
                        selected = getattr(dm_run, "selected_upload_id", None) or (dm_run.get("selected_upload_id") if isinstance(dm_run, dict) else None)
                        print(f"🏁 Auto DecisionMaker: {status} (chapter_id={chapter_id}, selected_upload_id={selected})")

                except Exception as e:
                    print(f"[WARN] Auto DecisionMaker trigger failed: {e}")


            try:
                bg_loop.run_until_complete(runner())
            except Exception as e:
                print("Error in evaluation background thread:", repr(e))
                traceback.print_exception(type(e), e, e.__traceback__)
            finally:
                bg_loop.close()

        threading.Thread(target=run_graph_background, daemon=True).start()

        return result



def increment_progress(contributor_id, chapter_id, file_type):
    progress, _ = ChapterContributionProgress.objects.get_or_create(
        contributor_id=contributor_id,
        chapter_id=chapter_id
    )

    if file_type == "pdf":
        progress.pdf_count += 1
    elif file_type == "video":
        progress.video_count += 1
    elif file_type == "draft":
        progress.draft_count += 1

    progress.has_any_upload = True
    progress.save()

@csrf_exempt
def confirm_submission(request):
    contributor_id = request.session.get("contributor_id")
    course_id = request.session.get("course_id")
    chapter_id = request.session.get("chapter_id")

    if not all([contributor_id, course_id, chapter_id]):
        return JsonResponse({"error": "Missing session data"}, status=400)

    chapter = Chapter.objects.get(id=chapter_id)
    chapter_number = chapter.chapter_number

    # -------------------------------
    # Google Drive (OOP)
    # -------------------------------
    service = GoogleDriveAuthService.get_service()
    folder_service = GoogleDriveFolderService(service)

    oer_root_id = folder_service.get_or_create_folder("oer_content")

    # Root folders
    pdf_root_id = folder_service.get_or_create_folder(
        settings.GOOGLE_DRIVE_FOLDERS["pdf"], oer_root_id
    )
    video_root_id = folder_service.get_or_create_folder(
        settings.GOOGLE_DRIVE_FOLDERS["videos"], oer_root_id
    )

    base_folder = f"{contributor_id}_{course_id}_{chapter_number}"

    # Contributor folders
    pdf_folder_id = folder_service.get_or_create_folder(base_folder, pdf_root_id)
    video_folder_id = folder_service.get_or_create_folder(base_folder, video_root_id)

    # -------------------------------
    # REQUIRED LangGraph State
    # -------------------------------
    state = {
        "contributor_id": contributor_id,
        "course_id": course_id,
        "chapter_id": chapter_id,
        "chapter_name": chapter.chapter_name,
        "drive_folders": {
            "pdf": pdf_folder_id,
            "videos": video_folder_id,
        },
    }

    print("Reached till submission orchestrator")

    # -------------------------------
    # Run Submission + Evaluation
    # -------------------------------
    result = SubmissionOrchestrator.submit_and_evaluate(state)

    if result.get("status") == "success":
        contributor = User.objects.get(id=contributor_id)

        ContributionSuccessEmail(
            contributor.email,
            contributor.first_name,
            chapter.course.course_name,
            chapter.chapter_name
        ).send()

        return render(request, "contributor/final_submission.html")

    return JsonResponse({"error": "Submission failed"}, status=500)


def contributor_upload_file(request):
    course_id = request.GET.get("course_id")
    chapter_id = request.GET.get("chapter_id")
    topic = unquote(request.GET.get('topic', ''))

    print("Upload file topic name: ", topic)

    tab = request.GET.get("tab", "content")

    if not all([course_id, chapter_id]):
        return HttpResponseBadRequest("Missing parameters")

    course = get_object_or_404(Course, id=course_id)
    chapter = get_object_or_404(Chapter, id=chapter_id)
    contributor_id = request.user.id

    if ContributorSubmissionService.has_existing_submission(contributor_id, chapter_id):
        return render(request, "contributor/after_submission.html")

    ContributorSessionService.store_submission_context(
        request,
        contributor_id=contributor_id,
        course_id=course_id,
        chapter_id=chapter_id,
        chapter_number=chapter.chapter_number,
        chapter_name=chapter.chapter_name,
        topic=topic
    )

    files = []

    try:
        # 🔹 Initialize Drive service (NEW AUTH CLASS)
        service = GoogleDriveAuthService.get_service()
        folder_service = GoogleDriveFolderService(service)

        oer_root_id = folder_service.get_or_create_folder("oer_content")

        # --- SAME LOGIC AS YOUR WORKING VERSION ---
        def get_files_from_folder(folder_type):
            folder_name = f"{contributor_id}_{course.id}_{chapter.chapter_number}"

            root_folder_id = folder_service.get_or_create_folder(
                settings.GOOGLE_DRIVE_FOLDERS[folder_type],
                oer_root_id
            )

            # Contributor chapter folder
            query = (
                "mimeType='application/vnd.google-apps.folder' "
                f"and name='{folder_name}' "
                f"and '{root_folder_id}' in parents "
                "and trashed=false"
            )

            folders = (
                service.files()
                .list(q=query, fields="files(id, name)")
                .execute()
                .get("files", [])
            )

            if not folders:
                return []

            chapter_folder_id = folders[0]["id"]
            collected_files = []

            # Fetch topic folders
            topic_folders = (
                service.files()
                .list(
                    q=(
                        "mimeType='application/vnd.google-apps.folder' "
                        f"and '{chapter_folder_id}' in parents "
                        "and trashed=false"
                    ),
                    fields="files(id, name)"
                )
                .execute()
                .get("files", [])
            )

            for topic_folder in topic_folders:
                topic_name = topic_folder["name"]

                # CRITICAL FILTER
                if topic_name != topic:
                    continue

                # print(f"📂 MATCHED TOPIC FOLDER: {topic_name}")

                topic_id = topic_folder["id"]

                files_result = (
                    service.files()
                    .list(
                        q=f"'{topic_id}' in parents and trashed=false",
                        fields="files(id, name, mimeType)"
                    )
                    .execute()
                )

                for f in files_result.get("files", []):
                    print(f"📄 FILE FOUND: {f['name']} (inside {topic_name})")

                    collected_files.append({
                        "id": f["id"],
                        "name": f["name"],
                        "mimeType": f["mimeType"],
                        "type": folder_type,
                        "topic": topic_name,
                    })

            return collected_files

        # Aggregate files
        for folder_type in ["drafts", "pdf", "videos", "assessments"]:
            files.extend(get_files_from_folder(folder_type))

    except RefreshError as e:
        print("[ERROR] Google token invalid:", e)

        token_path = settings.GOOGLE_TOKEN_FILE
        if os.path.exists(token_path):
            os.remove(token_path)

        messages.error(
            request,
            "⚠️ Your Google Drive session has expired. Please reconnect."
        )
        return redirect("contributor_dashboard")

    except Exception as e:
        print("[ERROR] Unexpected Drive issue:", e)
        messages.error(request, f"An unexpected error occurred: {e}")
        files = []

    # ---- Topic extraction ----
    raw_desc = chapter.description or ""
    topics = [t.strip() for t in re.split(r"[;,.]", raw_desc) if t.strip()]

    pdf_files = [f for f in files if f.get("type") == "pdf"]

    # ------------------------------
    # NEW: Assessments for this context
    # ------------------------------
    assessments = Assessment.objects.filter(
        contributor_id=request.user,
        course=course,
        chapter=chapter
    ).order_by("-id")

    context = {
        "course": course,
        "chapter": chapter,
        "files": files,
        "topic": topic,
        "topics": topics,
        "has_pdfs": len(pdf_files) > 0,
        "tab": tab,
        "assessments": assessments
    }

    return render(request, "contributor/contributor_upload_file.html", context)


@csrf_exempt
def upload_files(request):
    if request.method != "POST":
        return HttpResponseBadRequest("Invalid request")

    contributor_id = request.session.get("contributor_id")
    course_id = request.POST.get("course_id")
    chapter_id = request.POST.get("chapter_id")
    chapter_number = request.session.get('chapter_number')
    topic = request.POST.get("topic", "").strip()

    base_folder = f"{contributor_id}_{course_id}_{chapter_number}"
    drive = ContributorDriveUploadService()

    files = request.FILES.getlist("supporting_files")

    MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB

    for uploaded_file in files:
        if uploaded_file.size > MAX_FILE_SIZE:
            messages.error(
                request,
                f"{uploaded_file.name} exceeds 20MB limit"
            )
            continue

        # READ FILE CONTENT SAFELY
        file_bytes = uploaded_file.read()
        uploaded_file.seek(0)  # reset pointer (important)

        content_type = uploaded_file.content_type

        if content_type == "application/pdf":
            folder = drive.ensure_topic_folder(base_folder, "pdf", topic)
            file_type = "pdf"

        elif content_type.startswith("video/"):
            folder = drive.ensure_topic_folder(base_folder, "videos", topic)
            file_type = "video"

        elif uploaded_file.name.endswith((".doc", ".docx")):
            folder = drive.ensure_topic_folder(base_folder, "drafts", topic)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            file_type = "draft"

        else:
            continue

        # PASS BYTES, NOT FILE HANDLE
        drive.upload_file_bytes(
            file_bytes=file_bytes,
            filename=uploaded_file.name,
            folder_id=folder,
            content_type=content_type,
            contributor_id=contributor_id
        )

        # DB Progress update
        increment_progress(contributor_id, chapter_id, file_type)

    messages.success(request, "Files uploaded successfully")

    return redirect(
        f"{reverse('contributor_upload_file')}?"
        f"course_id={course_id}&"
        f"chapter_id={request.session.get('chapter_id')}&"
        f"topic={topic or ''}"
    )


# ---------------- EDITOR / DRAFT ---------------- #
@csrf_exempt
def contributor_editor(request):
    """Save drafts as DOCX or final submissions as PDF in Google Drive (OOP version)."""

    # --- Init Drive Services ---
    service = GoogleDriveAuthService.get_service()
    folder_service = GoogleDriveFolderService(service)

    # --- Session Data ---
    contributor_id = request.session.get("contributor_id", 101)
    course_id = request.session.get("course_id")
    chapter_number = request.session.get("chapter_number")
    chapter_name = request.session.get("chapter_name", "structured_query_language")
    topic_name = request.POST.get("topic") or request.GET.get("topic")

    print("Editor topic name: ", topic_name)

    if topic_name:
        topic_name = topic_name.replace("/", "_").strip()

    # --- Root folders ---
    oer_root_id = folder_service.get_or_create_folder("oer_content")

    drafts_root_id = folder_service.get_or_create_folder(
        settings.GOOGLE_DRIVE_FOLDERS["drafts"],
        oer_root_id
    )

    pdf_root_id = folder_service.get_or_create_folder(
        settings.GOOGLE_DRIVE_FOLDERS["pdf"],
        oer_root_id
    )

    # --- Contributor-level folders ---
    base_folder_name = f"{contributor_id}_{course_id}_{chapter_number}"

    drafts_folder_id = folder_service.get_or_create_folder(
        base_folder_name,
        drafts_root_id
    )

    pdf_folder_id = folder_service.get_or_create_folder(
        base_folder_name,
        pdf_root_id
    )

    # --- Topic-level folders ---
    if topic_name:
        drafts_topic_folder_id = folder_service.get_or_create_folder(
            topic_name,
            drafts_folder_id
        )
        pdf_topic_folder_id = folder_service.get_or_create_folder(
            topic_name,
            pdf_folder_id
        )
    else:
        drafts_topic_folder_id = drafts_folder_id
        pdf_topic_folder_id = pdf_folder_id

    # ======================================================
    # POST: Save Draft or Submit Draft
    # ======================================================
    if request.method == "POST":
        action = request.POST.get("action")  # 'draft' | 'submitDraft'
        content = request.POST.get("notes", "")
        filename = request.POST.get("filename", "draft")
        file_id = request.POST.get("file_id")

        try:
            # ---------- SAVE DRAFT ----------
            if action == "draft":
                from bs4 import BeautifulSoup

                doc = Document()
                soup = BeautifulSoup(content, "html.parser")

                for block in soup.find_all(["p", "div"]):
                    text = block.get_text(strip=True)
                    if text:
                        doc.add_paragraph(text)

                file_io = io.BytesIO()
                doc.save(file_io)
                increment_progress(contributor_id, request.session.get("chapter_id"), "draft")
                file_io.seek(0)

                media = MediaIoBaseUpload(
                    file_io,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    resumable=True
                )

                doc_filename = filename if filename.lower().endswith(".docx") else f"{filename}.docx"
                drive_filename = f"{doc_filename}"

                if file_id:
                    service.files().update(
                        fileId=file_id,
                        media_body=media
                    ).execute()
                else:
                    service.files().create(
                        body={
                            "name": drive_filename,
                            "parents": [drafts_topic_folder_id]
                        },
                        media_body=media,
                        fields="id"
                    ).execute()

            # ---------- SUBMIT AS PDF ----------
            elif action == "submitDraft":
                pdf_io = io.BytesIO()
                result = pisa.CreatePDF(io.StringIO(content), dest=pdf_io)

                if result.err:
                    raise Exception("PDF generation failed")

                pdf_io.seek(0)

                MAX_EDITOR_SIZE = 40 * 1024 * 1024  # 40MB

                pdf_size = len(pdf_io.getvalue())

                if pdf_size > MAX_EDITOR_SIZE:
                    messages.error(
                        request,
                        "Generated PDF exceeds 40MB limit."
                    )
                    return redirect(request.path)

                media = MediaIoBaseUpload(
                    pdf_io,
                    mimetype="application/pdf",
                    resumable=True
                )

                pdf_filename = filename if filename.lower().endswith(".pdf") else f"{filename}.pdf"
                if not pdf_filename.startswith(f"{contributor_id}_"):
                    pdf_filename = f"{pdf_filename}"

                service.files().create(
                    body={
                        "name": pdf_filename,
                        "parents": [pdf_topic_folder_id]
                    },
                    media_body=media,
                    fields="id"
                ).execute()

                increment_progress(contributor_id, request.session.get("chapter_id"), "pdf")

                # Delete old draft if editing existing
                if file_id:
                    try:
                        service.files().delete(fileId=file_id).execute()
                    except Exception as e:
                        print(f"[WARN] Could not delete draft {file_id}: {e}")

        except Exception as e:
            print(f"[ERROR] Draft action failed: {e}")
            messages.error(request, f"Failed to save draft: {e}")

    # ======================================================
    # Fetch existing drafts (for UI)
    # ======================================================
    try:
        results = service.files().list(
            q=f"'{drafts_folder_id}' in parents and trashed=false",
            fields="files(id, name, createdTime)"
        ).execute()
        files = results.get("files", [])
    except Exception as e:
        print(f"[ERROR] Failed to fetch drafts: {e}")
        files = []

    # --- Restore GET params and redirect ---
    request.GET = request.GET.copy()
    request.GET["course_id"] = str(course_id)
    request.GET["chapter_id"] = str(request.session.get("chapter_id"))
    request.GET["topic"] = topic_name or ""

    return redirect(
        f"{reverse('contributor_upload_file')}?"
        f"course_id={course_id}&"
        f"chapter_id={request.session.get('chapter_id')}&"
        f"topic={topic_name or ''}"
    )



# ---------------- LOAD FILE CONTENT ---------------- #
@csrf_exempt
def load_file(request):
    print("Load file")
    service = GoogleDriveAuthService.get_service()
    file_id = request.GET.get('file_id')

    print(file_id)

    if not file_id:
        return JsonResponse({'error': 'file_id is required'}, status=400)

    try:
        request_file = service.files().get(fileId=file_id, fields='mimeType').execute()
        mime_type = request_file.get('mimeType', 'text/html')

        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, service.files().get_media(fileId=file_id))
        done = False
        while not done:
            _, done = downloader.next_chunk()
        fh.seek(0)

        print("Loading file_id:", file_id)
        print("Service object:", service)

        if 'text/html' in mime_type:
            content = fh.getvalue().decode('utf-8')
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            import docx
            doc = docx.Document(fh)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Join with <p> tags for TinyMCE
            content = ''.join(f'<p>{p}</p>' for p in paragraphs)
        else:
            content = f"<p>Cannot edit file of type {mime_type} in the editor.</p>"

        return JsonResponse({'content': content})

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
def delete_drive_file(request):
    """Delete a Google Drive file permanently."""
    if request.method == 'POST':
        file_id = request.POST.get('file_id')
        service = GoogleDriveAuthService.get_service()
        try:
            service.files().delete(fileId=file_id).execute()
            return JsonResponse({'success': True, 'message': 'File deleted successfully.'})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})
    # Redirect back to original submission page
    course_id = request.session.get("course_id")
    chapter_id = request.session.get("chapter_id")
    return redirect(f'/dashboard/contributor/submit_content/?course_id={course_id}&chapter_id={chapter_id}')
    # return JsonResponse({'success': False, 'message': 'Invalid request'})


# @csrf_exempt
# def submit_assessment(request):
#     course_id = request.session.get("course_id")
#     chapter_id = request.session.get("chapter_id")
#
#     if not course_id or not chapter_id:
#         messages.error(request, "Course or Chapter not found in session.")
#         return redirect("/dashboard/contributor/submit_content/")
#
#     course = Course.objects.get(id=course_id)
#     chapter = Chapter.objects.get(id=chapter_id)
#
#     if request.method == 'POST':
#         # Extract all questions dynamically
#         questions_data = []
#         i = 0
#         while f'questions[{i}][question]' in request.POST:
#             q_text = request.POST[f'questions[{i}][question]']
#             correct = int(request.POST[f'questions[{i}][correct]'])
#             options = request.POST.getlist(f'questions[{i}][options][]')
#             questions_data.append({'text': q_text, 'correct': correct, 'options': options})
#             i += 1
#
#         assessment = Assessment.objects.create(course=course, chapter=chapter, contributor_id=request.user)
#
#         for q in questions_data:
#             question = Question.objects.create(
#                 assessment=assessment,
#                 text=q['text'],
#                 correct_option=q['correct']
#             )
#             for opt_text in q['options']:
#                 Option.objects.create(question=question, text=opt_text)
#
#     return redirect(f'/dashboard/contributor/submit_content/?course_id={course_id}&chapter_id={chapter_id}')


@csrf_exempt
def gemini_chat(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    try:
        # multipart/form-data
        user_message = request.POST.get("message", "").strip()
        files = request.FILES.getlist("files")  # multiple files

        if not user_message and not files:
            return JsonResponse({"error": "Message or files required"}, status=400)

        content = []

        # add text first
        if user_message:
            content.append({"type": "text", "text": user_message})

        # add files (works best for images)
        for f in files:
            file_bytes = f.read()
            mime_type = f.content_type or "application/octet-stream"

            content.append({
                "type": "media",
                "mime_type": mime_type,
                "data": file_bytes
            })

        # send to Gemini via LangChain
        response = llm.invoke([HumanMessage(content=content)])

        return JsonResponse({"reply": response.content})

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)


@csrf_exempt
def generate_assessment(request):
    """Generate MCQ assessment from selected PDFs using Gemini."""

    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    try:
        # =====================================================
        # 1. BASIC VALIDATION
        # =====================================================
        course_id = request.POST.get("course_id")
        chapter_id = request.POST.get("chapter_id")
        topic_name = request.POST.get("topic")

        difficulty = request.POST.get("difficulty", "Medium")
        question_count = int(request.POST.get("question_count", 10))
        custom_prompt = request.POST.get("custom_prompt", "").strip()

        if not all([course_id, chapter_id, topic_name]):
            return JsonResponse(
                {"error": "Missing required fields"},
                status=400
            )

        contributor_id = request.session.get("contributor_id")
        if not contributor_id:
            return JsonResponse(
                {"error": "Contributor not found in session"},
                status=403
            )

        selected_file_ids = request.POST.getlist("selected_files")

        if not selected_file_ids:
            return JsonResponse(
                {"error": "No PDF files selected."},
                status=400
            )

        # =====================================================
        # 2. PREVENT PDF REUSE
        # =====================================================
        already_used = AssessmentSource.objects.filter(
            assessment__contributor_id=contributor_id,
            drive_file_id__in=selected_file_ids
        ).values_list("drive_file_id", flat=True)

        if already_used:
            service = GoogleDriveAuthService.get_service()

            used_ids = AssessmentSource.objects.filter(
                assessment__contributor_id=contributor_id,
                drive_file_id__in=selected_file_ids
            ).values_list("drive_file_id", flat=True)

            used_files = []

            for fid in used_ids:
                try:
                    meta = service.files().get(
                        fileId=fid,
                        fields="name"
                    ).execute()

                    used_files.append(meta.get("name", "Unknown File"))

                except Exception:
                    used_files.append("Unknown File")

            return JsonResponse({
                "error": "Some selected PDFs were already used: "
                         + ", ".join(used_files)
            }, status=400)

        # =====================================================
        # 3. GOOGLE DRIVE DOWNLOAD
        # =====================================================
        service = GoogleDriveAuthService.get_service()
        pdf_texts = []

        for file_id in selected_file_ids:
            file_io = io.BytesIO()

            downloader = MediaIoBaseDownload(
                file_io,
                service.files().get_media(fileId=file_id)
            )

            done = False
            while not done:
                _, done = downloader.next_chunk()

            file_io.seek(0)

            try:
                reader = PdfReader(file_io)
                text = "\n".join(
                    page.extract_text() or ""
                    for page in reader.pages
                )

                if text.strip():
                    pdf_texts.append(text)

            except Exception as e:
                print(f"[WARN] Skipping PDF {file_id}: {e}")

        if not pdf_texts:
            return JsonResponse(
                {"error": "No readable text found."},
                status=400
            )

        combined_text = "\n".join(pdf_texts)[:15000]

        # =====================================================
        # 4. FETCH EXISTING QUESTIONS (UNIQUENESS)
        # =====================================================
        existing_questions = Question.objects.filter(
            assessment__contributor_id=contributor_id
        ).values_list("text", flat=True)[:200]

        existing_questions_text = "\n".join(existing_questions)

        # =====================================================
        # 5. GEMINI PROMPT
        # =====================================================
        prompt = f"""
You are an expert educational assessment generator.

Generate {question_count} UNIQUE multiple-choice questions.

STRICT RULES:
- Difficulty Level: {difficulty}
- Topic: {topic_name}
- Questions must NOT repeat or paraphrase previous ones.
- Focus on conceptual understanding.
- Exactly 4 options.
- Only ONE correct option.
- Return ONLY valid JSON.

Additional Instructions:
{custom_prompt}

Previously generated questions
(Do NOT repeat or rephrase):

{existing_questions_text}

Content Source:
{combined_text}

Return JSON:
{{
  "questions":[
    {{
      "text":"...",
      "options":["A","B","C","D"],
      "correct_option":1
    }}
  ]
}}
"""

        response = llm.invoke(prompt)
        response_text = response.content.strip()

        cleaned_text = re.sub(
            r"^```(?:json)?\s*|\s*```$",
            "",
            response_text,
            flags=re.DOTALL
        ).strip()

        try:
            result = json.loads(cleaned_text)
        except json.JSONDecodeError:
            return JsonResponse({
                "error": "Gemini returned invalid JSON",
                "raw": cleaned_text
            }, status=500)

        # =====================================================
        # 6. CREATE ASSESSMENT
        # =====================================================
        contributor = User.objects.get(id=contributor_id)
        course = Course.objects.get(id=course_id)
        chapter = Chapter.objects.get(id=chapter_id)

        assessment = Assessment.objects.create(
            course=course,
            chapter=chapter,
            contributor_id=contributor,
            topic=topic_name
        )

        # =====================================================
        # 7. SAVE USED PDF SOURCES (WITH REAL FILENAMES)
        # =====================================================

        file_metadata_map = {}

        # Fetch filenames from Google Drive
        for file_id in selected_file_ids:
            try:
                meta = service.files().get(
                    fileId=file_id,
                    fields="id,name"
                ).execute()

                file_metadata_map[file_id] = meta.get(
                    "name",
                    "Unknown File"
                )

            except Exception as e:
                print(f"[WARN] Could not fetch name for {file_id}: {e}")
                file_metadata_map[file_id] = "Unknown File"


        # Save sources
        for file_id in selected_file_ids:
            AssessmentSource.objects.create(
                assessment=assessment,
                drive_file_id=file_id,
                file_name=file_metadata_map.get(
                    file_id,
                    "Unknown File"
                )
            )
        # =====================================================
        # 8. HARD DUPLICATE PROTECTION
        # =====================================================
        existing_set = set(
            Question.objects.filter(
                assessment__contributor_id=contributor_id
            ).values_list("text", flat=True)
        )

        saved_count = 0

        for q_data in result.get("questions", []):

            q_text = q_data.get("text", "").strip()

            if not q_text or q_text in existing_set:
                continue

            question = Question.objects.create(
                assessment=assessment,
                text=q_text,
                correct_option=q_data.get(
                    "correct_option", 0
                )
            )

            existing_set.add(q_text)
            saved_count += 1

            for opt in q_data.get("options", []):
                Option.objects.create(
                    question=question,
                    text=opt
                )

        if saved_count == 0:
            assessment.delete()
            return JsonResponse({
                "error":
                    "All generated questions were duplicates."
            }, status=400)

        # =====================================================
        # 9. REDIRECT
        # =====================================================
        from django.urls import reverse

        return JsonResponse({
            "success": True,
            "redirect_url": reverse(
                "generated_assessment_form",
                args=[assessment.id]
            )
        })

    except Exception as e:
        print("[ERROR] Assessment generation failed:", e)
        return JsonResponse(
            {"error": str(e)},
            status=500
        )


def generated_assessment_form(request, assessment_id):
    """Display generated questions and options."""
    assessment = get_object_or_404(Assessment, id=assessment_id)
    questions = assessment.questions.prefetch_related('options').all()

    return render(request, "contributor/generated_assessment.html", {
        "assessment": assessment,
        "questions": questions,
    })


def after_submission(request):
    print("After submission view called")
    # generate_expertise()
    # Clear all session data safely
    return render(request, 'contributor/after_submission.html')



# External Resources
def list_resources(request):
    course_id = request.GET.get("course_id")
    chapter_id = request.GET.get("chapter_id")
    topic = request.GET.get("topic")

    qs = ExternalResource.objects.filter(
        course_id=course_id,
        chapter_id=chapter_id,
        topic=topic
    ).order_by("-created_at")

    data = [{
        "id": r.id,
        "title": r.title,
        "url": r.url,
        "type": r.resource_type
    } for r in qs]

    return JsonResponse({"resources": data})


@csrf_exempt
def add_resource(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    title = request.POST.get("title")
    url = request.POST.get("url")
    resource_type = request.POST.get("type", "youtube")

    course_id = request.POST.get("course_id")
    chapter_id = request.POST.get("chapter_id")
    topic = request.POST.get("topic")

    if not title or not url:
        return JsonResponse({"error": "Title and URL required"}, status=400)

    r = ExternalResource.objects.create(
        course_id=course_id,
        chapter_id=chapter_id,
        topic=topic,
        title=title,
        url=url,
        resource_type=resource_type,
        created_by=request.user
    )

    return JsonResponse({"success": True, "id": r.id})


@csrf_exempt
def delete_resource(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    rid = request.POST.get("id")
    ExternalResource.objects.filter(id=rid, created_by=request.user).delete()

    return JsonResponse({"success": True})


def run_auto_submit(contributor_id, chapter_id):
    chapter = Chapter.objects.get(id=chapter_id)
    course_id = chapter.course.id
    chapter_number = chapter.chapter_number

    service = GoogleDriveAuthService.get_service()
    folder_service = GoogleDriveFolderService(service)

    oer_root_id = folder_service.get_or_create_folder("oer_content")

    pdf_root_id = folder_service.get_or_create_folder(
        settings.GOOGLE_DRIVE_FOLDERS["pdf"], oer_root_id
    )
    video_root_id = folder_service.get_or_create_folder(
        settings.GOOGLE_DRIVE_FOLDERS["videos"], oer_root_id
    )

    base_folder = f"{contributor_id}_{course_id}_{chapter_number}"

    pdf_folder_id = folder_service.get_or_create_folder(base_folder, pdf_root_id)
    video_folder_id = folder_service.get_or_create_folder(base_folder, video_root_id)

    state = {
        "contributor_id": contributor_id,
        "course_id": course_id,
        "chapter_id": chapter_id,
        "chapter_name": chapter.chapter_name,
        "drive_folders": {
            "pdf": pdf_folder_id,
            "videos": video_folder_id,
        },
    }

    result = SubmissionOrchestrator.submit_and_evaluate(state)

    return result


def auto_submit_expired_deadlines():
    expired = ChapterPolicy.objects.filter(
        current_deadline__isnull=False,
        current_deadline__lt=timezone.now()
    )

    for policy in expired:
        chapter = policy.chapter

        progresses = ChapterContributionProgress.objects.filter(
            chapter=chapter,
            has_any_upload=True,
            auto_submitted=False
        )

        for progress in progresses:
            # call same pipeline you use in confirm_submission
            run_auto_submit(progress.contributor_id, chapter.id)

            progress.auto_submitted = True
            progress.save()

import json
from django.http import JsonResponse

from langgraph_agents.review_graph.review_workflow import (
    compiled_review_graph
)


# async def check_topic_quality(request):
#
#     data = json.loads(request.body)
#
#     topic = data.get("topic")
#     notes = data.get("notes")
#     files = data.get("files", [])
#     level = data.get("target_level", "undergrad")
#
#     result = await compiled_review_graph.ainvoke({
#         "topic": topic,
#         "notes": notes,
#         "files": files,
#         "target_level": level
#     })
#
#     clarity = result.get("clarity_review", {})
#     engagement = result.get("engagement_review", {})
#
#     suggestions = (
#             clarity.get("suggestions", [])
#             + engagement.get("suggestions", [])
#     )
#
#     return JsonResponse({
#         # "clarity_score": clarity.get("clarity_score"),
#         # "engagement_score": engagement.get("engagement_score"),
#         "suggestions": suggestions
#     })


async def check_topic_quality(request):

    data = json.loads(request.body)

    topic = data.get("topic")
    notes = data.get("notes")
    files = data.get("files", [])
    level = data.get("target_level", "undergrad")

    result = await run_review_pipeline({
        "topic": topic,
        "notes": notes,
        "files": files,
        "target_level": level
    })

    clarity = result.get("clarity_review", {})
    engagement = result.get("engagement_review", {})

    suggestions = (
            clarity.get("suggestions", [])
            + engagement.get("suggestions", [])
    )

    return JsonResponse({
        "suggestions": suggestions
    })